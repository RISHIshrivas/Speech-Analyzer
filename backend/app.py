import os
from io import BytesIO
from flask import Flask, request, jsonify, send_from_directory, send_file
from flask_cors import CORS

try:
    from docx import Document
except Exception:
    Document = None  # optional

BASE_DIR = os.path.dirname(__file__)
STATIC_DIR = os.path.abspath(os.path.join(BASE_DIR, "..", "static"))

app = Flask(__name__, static_folder=STATIC_DIR, static_url_path="/")
CORS(app)

# In-memory store (swap for a real DB in production)
speech_text = []


@app.get("/")
def index():
    """Serve the front-end index.html."""
    return send_from_directory(STATIC_DIR, "index.html")


@app.post("/save")
def save_speech():
    """Save recognized text chunk from the browser."""
    data = request.get_json(silent=True) or {}
    text = (data.get("text") or "").strip()
    if text:
        speech_text.append(text)
        return jsonify({"message": "saved", "total_entries": len(speech_text)})
    return jsonify({"message": "no text provided"}), 400


@app.get("/summary")
def get_summary():
    """Return a simple concatenation of all stored chunks."""
    summary_text = " ".join(speech_text)
    return jsonify({"summary": summary_text})


@app.post("/clear")
def clear_transcript():
    """Clear all stored speech text."""
    speech_text.clear()
    return jsonify({"message": "cleared", "total_entries": 0})


@app.get("/export/docx")
def export_docx():
    """Export stored transcript as a Word document."""
    if Document is None:
        return jsonify({"error": "python-docx not installed on server"}), 500
    if not speech_text:
        return jsonify({"error": "no text to export"}), 400

    doc = Document()
    doc.add_heading("Speech Analyzer Transcript", 0)
    for chunk in speech_text:
        doc.add_paragraph(chunk)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return send_file(
        buf,
        as_attachment=True,
        download_name="transcript.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


if __name__ == "__main__":
    import webbrowser
    port = int(os.environ.get("PORT", 5000))
    # Auto-open browser at localhost (mic works only here in dev)
    webbrowser.open(f"http://localhost:{port}")
    app.run(host="0.0.0.0", port=port, debug=True)
